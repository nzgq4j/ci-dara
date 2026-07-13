# modal/app.py

import modal
from fastapi import Depends, HTTPException, Request, status
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
import os
import re

# ── Image definition ──────────────────────────────────────────────────────────

image = (
    modal.Image.debian_slim(python_version="3.12")
    .pip_install([
        "pdfplumber==0.11.4",
        "python-docx==1.1.2",
        "spacy==3.8.14",
        "httpx",
        "fastapi[standard]",
    ])
    .run_commands(
        "python -m spacy download en_core_web_md"
    )
)

app = modal.App("dara-parser", image=image)

# ── Authentication ────────────────────────────────────────────────────────────

auth_scheme = HTTPBearer()

def verify_token(
    token: HTTPAuthorizationCredentials = Depends(auth_scheme),
):
    expected = os.environ.get("AUTH_TOKEN", "")
    if not expected or token.credentials != expected:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Unauthorized",
        )

# ── Domain entity patterns ────────────────────────────────────────────────────

FAR_PATTERN    = re.compile(r'\bFAR\s+\d{2}\.\d{3}-\d+\b', re.IGNORECASE)
DFARS_PATTERN  = re.compile(r'\bDFARS\s+\d{3}\.\d{3}-\d+\b', re.IGNORECASE)
DID_PATTERN    = re.compile(r'\bDI-[A-Z]+-\d+\b', re.IGNORECASE)
NIST_PATTERN   = re.compile(r'\bNIST\s+SP\s+\d{3}-\d+[A-Za-z]*\b', re.IGNORECASE)
MIL_PATTERN    = re.compile(r'\bMIL-STD-\d+[A-Z]*\b', re.IGNORECASE)

MODAL_VERBS = {
    "shall", "must", "will", "should",
    "may not", "must not",
    "is required to", "are required to",
    "is to be", "are to be",
}

CONDITIONAL_TRIGGERS = {
    "if ", "when ", "upon ", "unless ", "except ",
    "provided that", "in the event of", "in the event that",
    "only if", "subject to", "contingent upon", "wherever applicable",
}

CONDITION_TYPE_MAP = {
    "if ": "IF", "when ": "WHEN", "upon ": "UPON",
    "unless ": "UNLESS", "except ": "EXCEPT",
    "provided that": "IF", "in the event of": "UPON",
    "in the event that": "UPON", "only if": "IF",
    "subject to": "IF", "contingent upon": "UPON",
    "wherever applicable": "IF",
}

_flag_counter = 0

# ── Text cleaning ─────────────────────────────────────────────────────────────

# Soft hyphen (U+00AD) marks an optional hyphenation point. pdfplumber preserves it in the extracted
# text — typically at the line break where a word was split ("com­\npliance") — but an LLM
# transcribing the same text emits the JOINED word ("compliance"). A downstream verbatim comparison of
# LLM output against this raw text then mismatches on the artifact alone. Rejoin the word by removing
# the soft hyphen together with any whitespace (including a single line break) that follows it, then
# drop zero-width characters and the BOM. Lossless de-artifacting of the source text — no real content
# is altered.
_SOFT_HYPHEN_BREAK = re.compile(chr(0x00AD) + r"[ \t]*\n?[ \t]*")  # soft hyphen + trailing ws / one linebreak
_ZERO_WIDTH        = re.compile("[" + chr(0x200B) + chr(0x200C) + chr(0x200D) + chr(0xFEFF) + "]")


def clean_extracted_text(text: str) -> str:
    if not text:
        return text
    text = _SOFT_HYPHEN_BREAK.sub("", text)
    text = _ZERO_WIDTH.sub("", text)
    return text

# ── Main endpoint ─────────────────────────────────────────────────────────────

@app.function(
    cpu=2,
    memory=2048,
    timeout=180,
    secrets=[modal.Secret.from_name("dara-parser-secret")],
)
@modal.fastapi_endpoint(method="POST")
async def parse_document(
    request: Request,
    _: None = Depends(verify_token),
):
    import httpx
    import time

    body = await request.json()
    document_url = body.get("document_url")
    document_id  = body.get("document_id")
    doc_type     = body.get("doc_type", "pdf")

    if not document_url or not document_id:
        raise HTTPException(
            status_code=400,
            detail="document_url and document_id are required"
        )

    start_ms = int(time.time() * 1000)

    async with httpx.AsyncClient(timeout=30) as client:
        response = await client.get(document_url)
        response.raise_for_status()
        doc_bytes = response.content

    if doc_type == "pdf":
        result = _parse_pdf(doc_bytes, document_id)
    else:
        result = _parse_docx(doc_bytes, document_id)

    result["processing_time_ms"] = int(time.time() * 1000) - start_ms
    result["document_id"]        = document_id
    result["doc_type"]           = doc_type
    result["schema_version"]     = "1.0"
    result["parser_version"]     = "0.1.0"

    return result

# ── Helper functions ──────────────────────────────────────────────────────────

def _detect_modal_verbs(text: str) -> list[str]:
    text_lower = text.lower()
    return [v for v in MODAL_VERBS if v in text_lower]


def _is_passive(sent) -> bool:
    return any(tok.dep_ in ("nsubjpass", "auxpass") for tok in sent)


def _has_conditional_trigger(text: str) -> bool:
    text_lower = text.lower()
    return any(t in text_lower for t in CONDITIONAL_TRIGGERS)


def _extract_conditional_trigger(text: str, sentence_id: str) -> dict | None:
    text_lower = text.lower()
    for trigger, ctype in CONDITION_TYPE_MAP.items():
        if trigger in text_lower:
            idx = text_lower.index(trigger)
            return {
                "trigger_id":           f"cond-{sentence_id}",
                "sentence_id":          sentence_id,
                "condition_type":       ctype,
                "trigger_text":         text[idx:idx + 80],
                "scope_text":           text,
                "condition_confidence": "EXPLICIT",
            }
    return None


def _extract_modal_candidate(sent, sent_id: str, para_id: str, section_id) -> dict:
    subject      = None
    verb_phrase  = None
    obj          = None
    modal_verb   = None

    for token in sent:
        if token.tag_ == "MD" and modal_verb is None:
            modal_verb = token.text.lower()
        if token.dep_ in ("nsubj", "nsubjpass") and subject is None:
            subject = token.text
        if token.dep_ in ("dobj", "attr") and obj is None:
            obj = token.text
        if token.dep_ == "ROOT" and verb_phrase is None:
            verb_phrase = token.text

    subject_inferred = subject is None
    if subject_inferred:
        subject = "Contractor"

    modal_verbs    = _detect_modal_verbs(sent.text)
    final_modal    = modal_verb or (modal_verbs[0] if modal_verbs else "shall")
    is_passive     = _is_passive(sent)

    if not subject_inferred and obj:
        svo_confidence = "HIGH"
    elif subject_inferred or not obj:
        svo_confidence = "MEDIUM"
    else:
        svo_confidence = "LOW"

    return {
        "candidate_id":          f"cand-{sent_id}",
        "sentence_id":           sent_id,
        "paragraph_id":          para_id,
        "section_id":            section_id,
        "source_text":           sent.text.strip(),
        "modal_verb":            final_modal,
        "modal_class":           "MANDATORY",
        "subject":               subject,
        "subject_inferred":      subject_inferred,
        "subject_confidence":    "LOW" if subject_inferred else "HIGH",
        "verb_phrase":           verb_phrase,
        "object":                obj,
        "is_passive":            is_passive,
        "svo_confidence":        svo_confidence,
        "section_context":       None,
        "parent_paragraph_text": None,
    }


def _detect_ibr_flags(text: str, sentence_id, table_row_id) -> list[dict]:
    global _flag_counter
    flags    = []
    patterns = [
        (FAR_PATTERN,   "FAR"),
        (DFARS_PATTERN, "DFARS"),
        (DID_PATTERN,   "DID"),
        (NIST_PATTERN,  "NIST"),
        (MIL_PATTERN,   "MIL_STD"),
    ]
    for pattern, citation_type in patterns:
        for match in pattern.finditer(text):
            _flag_counter += 1
            flags.append({
                "flag_id":         f"ibr-{_flag_counter}",
                "sentence_id":     sentence_id,
                "table_row_id":    table_row_id,
                "citation_text":   match.group(),
                "citation_type":   citation_type,
                "traversal_status": "PENDING",
            })
    return flags


def _process_sentences(nlp, text: str, para_id: str, section_id):
    """Run spaCy on a paragraph and return sentences, modal candidates,
    conditional triggers, named entities, and IbR flags."""
    sentences           = []
    modal_candidates    = []
    conditional_triggers = []
    named_entities      = []
    ibr_flags           = []

    doc = nlp(text)
    for sent_i, sent in enumerate(doc.sents):
        sent_text = sent.text.strip()
        if len(sent_text) < 10:
            continue

        sent_id      = f"sent-{para_id}-{sent_i}"
        has_modal    = bool(_detect_modal_verbs(sent_text))
        is_passive   = _is_passive(sent)
        is_cond      = _has_conditional_trigger(sent_text)

        sentences.append({
            "sentence_id":    sent_id,
            "paragraph_id":   para_id,
            "section_id":     section_id,
            "text":           sent_text,
            "char_start":     sent.start_char,
            "char_end":       sent.end_char,
            "has_modal_verb": has_modal,
            "is_passive":     is_passive,
            "is_conditional": is_cond,
        })

        if has_modal:
            modal_candidates.append(
                _extract_modal_candidate(sent, sent_id, para_id, section_id)
            )

        if is_cond:
            trigger = _extract_conditional_trigger(sent_text, sent_id)
            if trigger:
                conditional_triggers.append(trigger)

        for ent in sent.ents:
            named_entities.append({
                "entity_id":   f"ent-{sent_id}-{ent.start}",
                "sentence_id": sent_id,
                "text":        ent.text,
                "label":       ent.label_,
                "start_char":  ent.start_char - sent.start_char,
                "end_char":    ent.end_char   - sent.start_char,
                "source":      "statistical",
            })

        ibr_flags.extend(_detect_ibr_flags(sent_text, sent_id, None))

    return sentences, modal_candidates, conditional_triggers, named_entities, ibr_flags


# ── PDF Parser ────────────────────────────────────────────────────────────────

def _parse_pdf(doc_bytes: bytes, document_id: str) -> dict:
    import pdfplumber
    import io
    import spacy

    nlp = spacy.load("en_core_web_md")

    pages                = []
    sections             = []
    tables               = []
    paragraphs           = []
    sentences            = []
    modal_candidates     = []
    conditional_triggers = []
    named_entities       = []
    ibr_flags            = []
    quality_gate_failures = []
    total_words          = 0

    with pdfplumber.open(io.BytesIO(doc_bytes)) as pdf:
        page_count = len(pdf.pages)

        for page_num, page in enumerate(pdf.pages, start=1):
            words      = page.extract_words(extra_attrs=["fontname", "size"])
            word_count = len(words)
            total_words += word_count
            image_only = word_count < 50

            if image_only:
                quality_gate_failures.append({
                    "gate":           "image_layer",
                    "affected_pages": [page_num],
                    "detail":         f"Page {page_num} has fewer than 50 words — likely image-only.",
                })

            pages.append({
                "page_number":    page_num,
                "width":          float(page.width),
                "height":         float(page.height),
                "word_count":     word_count,
                "has_text_layer": not image_only,
                "image_only":     image_only,
                "section_id":     None,
            })

            # Tables
            for tbl_i, table in enumerate(page.extract_tables()):
                if not table or len(table) < 2:
                    continue
                headers = [clean_extracted_text(str(c or "")).strip() for c in table[0]]
                rows    = []
                is_cdrl = any(
                    "cdrl" in h.lower() or "data item" in h.lower() or "did" in h.lower()
                    for h in headers
                )
                is_obligation_bearing = False

                for row_i, row in enumerate(table[1:]):
                    cells = {
                        headers[i]: clean_extracted_text(str(cell or "")).strip()
                        for i, cell in enumerate(row)
                        if i < len(headers)
                    }
                    reconstructed  = " | ".join(f"{k}: {v}" for k, v in cells.items() if v)
                    modal_found    = _detect_modal_verbs(reconstructed)
                    row_ibr        = _detect_ibr_flags(
                        reconstructed, None,
                        f"tbl-{page_num}-{tbl_i}-{row_i}"
                    )
                    if modal_found:
                        is_obligation_bearing = True
                    ibr_flags.extend(row_ibr)

                    rows.append({
                        "row_index":          row_i,
                        "cells":              cells,
                        "reconstructed_text": reconstructed,
                        "modal_verbs_found":  modal_found,
                        "ibr_flags":          [f["flag_id"] for f in row_ibr],
                    })

                tables.append({
                    "table_id":            f"tbl-{page_num}-{tbl_i}",
                    "page_number":         page_num,
                    "section_id":          None,
                    "bbox":                [0, 0, float(page.width), float(page.height)],
                    "headers":             headers,
                    "rows":                rows,
                    "is_cdrl":             is_cdrl,
                    "is_obligation_bearing": is_obligation_bearing,
                })

            # Body text
            page_text = clean_extracted_text(page.extract_text() or "")
            if not page_text.strip():
                continue

            para_id = f"para-p{page_num}"
            paragraphs.append({
                "paragraph_id":       para_id,
                "section_id":         None,
                "page_number":        page_num,
                "text":               page_text,
                "element_type":       "body",
                "list_level":         None,
                "parent_paragraph_id": None,
                "bbox":               None,
            })

            s, mc, ct, ne, ibr = _process_sentences(nlp, page_text, para_id, None)
            sentences.extend(s)
            modal_candidates.extend(mc)
            conditional_triggers.extend(ct)
            named_entities.extend(ne)
            ibr_flags.extend(ibr)

    quality_gate_passed = not any(
        f["gate"] == "structure_detection" for f in quality_gate_failures
    )

    return {
        "page_count":            page_count,
        "word_count":            total_words,
        "quality_gate_passed":   quality_gate_passed,
        "quality_gate_failures": quality_gate_failures,
        "pages":                 pages,
        "sections":              sections,
        "tables":                tables,
        "paragraphs":            paragraphs,
        "sentences":             sentences,
        "modal_candidates":      modal_candidates,
        "conditional_triggers":  conditional_triggers,
        "named_entities":        named_entities,
        "ibr_flags":             ibr_flags,
        "modal_candidate_count": len(modal_candidates),
        "table_count":           len(tables),
        "ibr_flag_count":        len(ibr_flags),
        "image_page_count":      sum(1 for p in pages if p["image_only"]),
    }


# ── DOCX Parser ───────────────────────────────────────────────────────────────

def _parse_docx(doc_bytes: bytes, document_id: str) -> dict:
    from docx import Document
    import io
    import spacy

    nlp = spacy.load("en_core_web_md")
    doc = Document(io.BytesIO(doc_bytes))

    sections             = []
    paragraphs           = []
    tables               = []
    sentences            = []
    modal_candidates     = []
    conditional_triggers = []
    named_entities       = []
    ibr_flags            = []
    para_index           = 0

    for para in doc.paragraphs:
        text = clean_extracted_text(para.text).strip()
        if not text:
            continue

        style_name    = para.style.name if para.style else ""
        is_heading    = style_name.startswith("Heading")
        heading_level = 0

        if is_heading:
            try:
                heading_level = int(style_name.split()[-1])
            except ValueError:
                heading_level = 1
            sections.append({
                "section_id":       f"sec-{para_index}",
                "heading_text":     text,
                "heading_level":    heading_level,
                "source_numbering": None,
                "synthetic_path":   f"S-{len(sections) + 1}",
                "page_start":       0,
                "page_end":         0,
                "parent_section_id": None,
                "font_size":        None,
                "bold":             True,
            })
            para_index += 1
            continue

        para_id = f"para-{para_index}"
        para_index += 1

        paragraphs.append({
            "paragraph_id":       para_id,
            "section_id":         None,
            "page_number":        None,
            "text":               text,
            "element_type":       "body",
            "list_level":         None,
            "parent_paragraph_id": None,
            "bbox":               None,
        })

        s, mc, ct, ne, ibr = _process_sentences(nlp, text, para_id, None)
        sentences.extend(s)
        modal_candidates.extend(mc)
        conditional_triggers.extend(ct)
        named_entities.extend(ne)
        ibr_flags.extend(ibr)

    # Tables
    for tbl_i, table in enumerate(doc.tables):
        if not table.rows:
            continue
        headers = [clean_extracted_text(cell.text).strip() for cell in table.rows[0].cells]
        rows    = []
        is_cdrl = any("cdrl" in h.lower() or "did" in h.lower() for h in headers)
        is_obligation_bearing = False

        for row_i, row in enumerate(table.rows[1:]):
            cells = {
                headers[i]: clean_extracted_text(cell.text).strip()
                for i, cell in enumerate(row.cells)
                if i < len(headers)
            }
            reconstructed = " | ".join(f"{k}: {v}" for k, v in cells.items() if v)
            modal_found   = _detect_modal_verbs(reconstructed)
            row_ibr       = _detect_ibr_flags(
                reconstructed, None, f"tbl-{tbl_i}-{row_i}"
            )
            if modal_found:
                is_obligation_bearing = True
            ibr_flags.extend(row_ibr)

            rows.append({
                "row_index":          row_i,
                "cells":              cells,
                "reconstructed_text": reconstructed,
                "modal_verbs_found":  modal_found,
                "ibr_flags":          [f["flag_id"] for f in row_ibr],
            })

        tables.append({
            "table_id":              f"tbl-{tbl_i}",
            "page_number":           None,
            "section_id":            None,
            "bbox":                  None,
            "headers":               headers,
            "rows":                  rows,
            "is_cdrl":               is_cdrl,
            "is_obligation_bearing": is_obligation_bearing,
        })

    total_words = sum(len(p["text"].split()) for p in paragraphs)

    return {
        "page_count":            None,
        "word_count":            total_words,
        "quality_gate_passed":   True,
        "quality_gate_failures": [],
        "pages":                 [],
        "sections":              sections,
        "tables":                tables,
        "paragraphs":            paragraphs,
        "sentences":             sentences,
        "modal_candidates":      modal_candidates,
        "conditional_triggers":  conditional_triggers,
        "named_entities":        named_entities,
        "ibr_flags":             ibr_flags,
        "modal_candidate_count": len(modal_candidates),
        "table_count":           len(tables),
        "ibr_flag_count":        len(ibr_flags),
        "image_page_count":      0,
    }